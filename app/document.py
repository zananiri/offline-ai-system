"""
Document conversion + OCR via Docling. Handles PDF, DOCX, PPTX, images, etc.
and returns clean Markdown, chunked for translation, or exports to Word.

Two separate OCR paths are used, deliberately:
- RapidOCR via Docling (default) — better accuracy on real-world scans,
  built-in table detection, handles skewed/photographed documents well.
  This is what every document uses unless told otherwise.
- Direct Tesseract, bypassing Docling (Hebrew only) — RapidOCR (and EasyOCR)
  have no Hebrew model at all, so Tesseract is the only engine here that
  can read Hebrew script. Docling's own Tesseract integration
  (TesseractCliOcrOptions) has a known, currently unresolved upstream bug:
  it writes the rendered page to a temp image file without embedding DPI
  metadata, so Tesseract can't determine the resolution, guesses a useless
  70 DPI, and aborts the page ("Too few characters. Skipping this page").
  See docling-project/docling-serve#282 and open-webui#15952 — this is not
  fixable via images_scale or any other documented Docling option. Instead,
  Hebrew documents are rendered to images ourselves (via pypdfium2) and OCR'd
  by calling Tesseract directly (via pytesseract) with an explicit DPI, which
  sidesteps the auto-detection failure entirely.

  Trade-off: the Hebrew path returns plain paragraph text per page, not
  Docling's structured Markdown (no table detection). Fine for translation
  and chat use; a Hebrew document with heavy tables will lose that structure.
"""
import html
import os
import re
import shutil
from pathlib import Path

import pypandoc
import pypdfium2 as pdfium
import pytesseract
from pytesseract import Output
from PIL import Image
from pypdf import PdfReader
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.pipeline_options import PdfPipelineOptions, RapidOcrOptions
from docling.datamodel.base_models import InputFormat

# Tesseract (the OCR *engine*, not the pytesseract Python wrapper) is a
# separate program that must be installed on the machine and discoverable
# on PATH. pytesseract just shells out to it. If it's not on PATH -- a very
# common setup gap on Windows, especially right after installing it without
# restarting the terminal -- pytesseract raises TesseractNotFoundError deep
# inside a subprocess call, which otherwise surfaces as a raw, unhelpful
# 500 traceback to the user. TESSERACT_CMD lets you point at an explicit
# install location if PATH still doesn't pick it up for some reason (e.g.
# installed for a different Windows user, or a portable install).
_TESSERACT_CMD_OVERRIDE = os.environ.get("TESSERACT_CMD")
if _TESSERACT_CMD_OVERRIDE:
    pytesseract.pytesseract.tesseract_cmd = _TESSERACT_CMD_OVERRIDE


def _tesseract_available() -> bool:
    cmd = pytesseract.pytesseract.tesseract_cmd
    # tesseract_cmd defaults to just "tesseract" (relies on PATH) unless
    # overridden above with an explicit full path.
    return shutil.which(cmd) is not None or Path(cmd).is_file()


class TesseractNotAvailableError(RuntimeError):
    """Raised when Hebrew OCR is requested but Tesseract isn't installed/found."""

# --- Default converter: RapidOCR via Docling (higher accuracy, no Hebrew support) ---
_default_pipeline_options = PdfPipelineOptions()
_default_pipeline_options.do_ocr = True
# Docling's default images_scale=1.0 renders pages at only 72 DPI, which is
# too low for reliable OCR. 3.0 -> ~216 DPI.
_default_pipeline_options.images_scale = 3.0
_default_pipeline_options.ocr_options = RapidOcrOptions()  # ONNX-based, light on CPU
_default_pipeline_options.do_table_structure = True

_default_converter = DocumentConverter(
    format_options={
        InputFormat.PDF: PdfFormatOption(pipeline_options=_default_pipeline_options),
    }
)

# --- Hebrew OCR: direct Tesseract call, bypassing Docling's OCR pipeline entirely ---
_HEBREW_OCR_LANG = "heb+eng"
_HEBREW_OCR_DPI = 300
_OCR_INPUT_EXTS = {".pdf", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}

# Tesseract confidence is 0-100 per recognized word. image_to_string (the
# old approach) returns Tesseract's best guess for EVERY detected region
# regardless of confidence -- on a noisy/low-quality scan that includes a
# lot of near-random character soup for regions Tesseract barely
# recognized at all (stray marks, table borders, watermark bleed-through,
# etc.), which is what shows up downstream as "unknown characters" in the
# extracted text. _ocr_image_filtered below uses image_to_data instead and
# drops anything under this threshold.
_OCR_MIN_CONFIDENCE = 40

# Forces the LSTM-only engine instead of Tesseract's default --oem 3
# (legacy + LSTM combined, which lets the older/less accurate legacy
# engine's output win in some cases). LSTM-only is consistently more
# accurate on Hebrew script in practice.
_HEBREW_OCR_OEM = 1

# Unicode bidi/formatting control characters: ZERO WIDTH SPACE / NON-JOINER /
# JOINER, LEFT-TO-RIGHT MARK, RIGHT-TO-LEFT MARK, the LTR/RTL/POP embedding
# & override marks, the newer directional isolates, and BOM.
_BIDI_CONTROL_CHARS_RE = re.compile("[\u200b-\u200f\u202a-\u202e\u2066-\u2069\ufeff]")


def strip_bidi_controls(text: str) -> str:
    """
    Strips invisible Unicode bidi/formatting control characters from OCR'd
    text.

    Tesseract's Hebrew OCR (lang="heb+eng") inserts these routinely to
    preserve correct reading order around embedded LTR runs (English words,
    dates, numbers) inside RTL text -- completely correct and invisible when
    the text is just *displayed*. But MADLAD-400's SentencePiece tokenizer
    treats them as ordinary characters, and when one lands mid-word it
    fragments tokenization badly enough that the model gives up translating
    and echoes the (still-Hebrew) input back unchanged instead. Observed in
    practice as a 100% "wrong language" rejection rate on Hebrew-sourced
    chunks. They carry no translatable meaning, so it's safe to drop them
    before the text is ever tokenized. See translate.py's _translate_once
    for the other half of this fix (it's applied there too, defensively).
    """
    if not text:
        return text
    return _BIDI_CONTROL_CHARS_RE.sub("", text)


# The Unicode replacement character, the object-replacement character, and
# the Private Use Area range -- these are what Tesseract (and some font
# encodings it misreads) emit when it "recognizes" a glyph shape but the
# result isn't a real character. They carry no translatable meaning and,
# like the bidi marks above, can fragment MADLAD-400's tokenization if left
# in. _OCR_MIN_CONFIDENCE filtering in _ocr_image_filtered catches most of
# this at the source, but this is a defensive second pass for whatever
# still slips through (e.g. a confidently-wrong glyph substitution).
_OCR_GARBAGE_CHARS_RE = re.compile("[\ufffc\ufffd\ue000-\uf8ff]")


def strip_ocr_noise(text: str) -> str:
    """Strips replacement/private-use-area characters left behind by OCR
    misreads. See _OCR_GARBAGE_CHARS_RE above for what this targets."""
    if not text:
        return text
    return _OCR_GARBAGE_CHARS_RE.sub("", text)


def _render_pdf_pages(file_path: str, dpi: int = _HEBREW_OCR_DPI) -> list[Image.Image]:
    pdf = pdfium.PdfDocument(file_path)
    scale = dpi / 72  # pypdfium2's render scale is relative to a 72-DPI baseline
    images = []
    for page in pdf:
        bitmap = page.render(scale=scale)
        pil_image = bitmap.to_pil()
        pil_image.info["dpi"] = (dpi, dpi)  # embed DPI so Tesseract doesn't need to guess
        images.append(pil_image)
    return images


# Below this many total characters, a PDF's "native text layer" is treated
# as not actually usable (e.g. a scanned PDF where only a stray bit of
# metadata happens to be real text) -- falls back to real OCR instead.
_MIN_NATIVE_TEXT_CHARS = 200


def _extract_native_pdf_text(file_path: str) -> str | None:
    """
    Extracts a PDF's own embedded text layer directly via pypdf -- no
    rendering or OCR involved at all. Many "PDFs" -- especially browser
    Print-to-PDF output, which is exactly what government/news pages
    saved as PDF usually are -- already contain a perfect, digitally
    authored text layer. Routing these through rasterize-then-Tesseract-OCR
    anyway is strictly worse: it throws away already-correct text and
    reintroduces real OCR misreads. Confirmed directly in practice: a
    Hebrew government-site PDF with a clean native text layer came out of
    the old rasterize+OCR path with fragments like "Nan?", stray Latin
    letters substituted for Hebrew words, and transposed digits, on a
    document that had zero errors in its own embedded text.

    Returns None (so the caller falls back to actual OCR) if there's no
    usable text layer at all (a genuinely scanned/image-only PDF), or if
    what's extracted doesn't look predominantly Hebrew -- the same
    corrupted-font-encoding failure mode _detect_script_is_hebrew_from_image
    already guards against (some Hebrew PDFs have a text layer where a
    meaningful share of glyphs decode to garbage/Latin-lookalike
    characters; that's not trustworthy to use verbatim either).
    """
    try:
        reader = PdfReader(file_path)
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception:
        return None
    if len(text) < _MIN_NATIVE_TEXT_CHARS:
        return None

    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return None
    hebrew_ratio = sum(1 for ch in letters if 0x0590 <= ord(ch) <= 0x05FF) / len(letters)
    if hebrew_ratio < _HEBREW_SCRIPT_RATIO_THRESHOLD:
        return None

    return text


def _ocr_image_filtered(img: Image.Image, dpi: int = _HEBREW_OCR_DPI) -> str:
    """
    Word-confidence-filtered OCR for a single image, replacing a plain
    pytesseract.image_to_string call.

    image_to_string keeps Tesseract's single best-guess glyph sequence for
    every detected text region no matter how low-confidence that guess
    was. image_to_data exposes per-word confidence, so this drops any word
    below _OCR_MIN_CONFIDENCE -- removing the noise that otherwise shows up
    as "unknown characters" -- while reconstructing paragraph/line breaks
    from Tesseract's block/paragraph/line numbering so downstream chunking
    (chunk_text, which splits on blank lines) still sees real paragraphs.
    """
    data = pytesseract.image_to_data(
        img, lang=_HEBREW_OCR_LANG,
        config=f"--psm 3 --oem {_HEBREW_OCR_OEM} --dpi {dpi}",
        output_type=Output.DICT,
    )

    lines: dict[tuple[int, int, int], list[str]] = {}
    for i, word in enumerate(data["text"]):
        word = word.strip()
        if not word:
            continue
        try:
            conf = float(data["conf"][i])
        except (TypeError, ValueError):
            conf = -1  # Tesseract uses -1 for non-text rows (e.g. block/line markers)
        if conf < _OCR_MIN_CONFIDENCE:
            continue
        key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        lines.setdefault(key, []).append(word)

    # data's rows are already in reading order, so sorting by the
    # (block, par, line) key preserves it -- group lines into paragraphs.
    paragraphs: dict[tuple[int, int], list[str]] = {}
    for (block, par, _line), words in sorted(lines.items()):
        paragraphs.setdefault((block, par), []).append(" ".join(words))

    return "\n\n".join(
        "\n".join(line_texts) for _, line_texts in sorted(paragraphs.items())
    ).strip()


def _ocr_hebrew(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        native_text = _extract_native_pdf_text(file_path)
        if native_text is not None:
            print(f"[ocr] {Path(file_path).name}: usable native PDF text layer found, "
                  "skipping rasterize+OCR entirely.")
            return strip_ocr_noise(strip_bidi_controls(html.unescape(native_text)))
        print(f"[ocr] {Path(file_path).name}: no usable native text layer, falling back to Tesseract OCR.")

    if not _tesseract_available():
        raise TesseractNotAvailableError(
            "Hebrew OCR requires Tesseract, which isn't installed or isn't on PATH. "
            "Install it from https://github.com/UB-Mannheim/tesseract/wiki, make sure "
            "to check the Hebrew language pack during setup, then restart the backend "
            "(PATH changes don't apply to already-running processes/terminals). "
            "If it's installed somewhere PATH doesn't cover, set the TESSERACT_CMD "
            "environment variable to the full path of tesseract.exe instead."
        )

    if suffix == ".pdf":
        images = _render_pdf_pages(file_path)
    else:
        img = Image.open(file_path)
        img.info["dpi"] = (_HEBREW_OCR_DPI, _HEBREW_OCR_DPI)
        images = [img]

    page_texts = []
    for img in images:
        # --psm 3: fully automatic page segmentation, explicitly WITHOUT an
        # orientation/script-detection (OSD) sub-pass — OSD is exactly the
        # sub-step that fails with bad DPI in Docling's wrapper, so we skip
        # it here by design. --dpi is passed explicitly as a second safeguard
        # on top of the DPI already embedded in the image itself. --oem 1
        # forces the LSTM-only engine (see _HEBREW_OCR_OEM). Confidence
        # filtering happens inside _ocr_image_filtered.
        text = _ocr_image_filtered(img, dpi=_HEBREW_OCR_DPI)
        if text:
            page_texts.append(text)

    joined = html.unescape("\n\n".join(page_texts))
    # Strip bidi/formatting marks and any remaining OCR-noise characters
    # right at the source, so every downstream consumer (translation, chat
    # context, summarization, the raw markdown returned to the UI) sees
    # clean text, not just the translation path.
    return strip_ocr_noise(strip_bidi_controls(joined))


# Threshold for _sample_text_hebrew_ratio / auto-detection below: if at
# least this fraction of the letters in a page's text are Hebrew-script,
# the document is treated as Hebrew. Deliberately low -- a Hebrew document
# with English proper nouns, numbers, or a mixed letterhead should still
# clear this easily, while a genuinely non-Hebrew document with a stray
# Hebrew word or two should not.
_HEBREW_SCRIPT_RATIO_THRESHOLD = 0.15


def _sample_text_hebrew_ratio(file_path: str, max_pages: int = 5) -> float | None:
    """
    Cheaply samples a PDF's existing text layer (no rendering/OCR at all)
    and returns the fraction of Hebrew-script letters among all letters
    found. Returns None if there's no extractable text layer at all (a
    fully scanned PDF), in which case _detect_script_is_hebrew falls back
    to image-based script detection instead.
    """
    try:
        reader = PdfReader(file_path)
    except Exception:
        return None
    text_parts = []
    for page in reader.pages[:max_pages]:
        try:
            text_parts.append(page.extract_text() or "")
        except Exception:
            continue
    sample = "".join(text_parts)
    letters = [ch for ch in sample if ch.isalpha()]
    if not letters:
        return None
    hebrew_letters = sum(1 for ch in letters if 0x0590 <= ord(ch) <= 0x05FF)
    return hebrew_letters / len(letters)


def _detect_script_is_hebrew_from_image(file_path: str) -> bool | None:
    """
    Renders (or opens) a single sample page/image and runs Tesseract's OSD
    (orientation + script detection) on it. This works directly off pixels,
    so -- unlike _sample_text_hebrew_ratio -- it can't be fooled by a PDF
    whose embedded text layer has a broken/non-standard font encoding
    (a real, observed failure mode: SOME glyphs in an otherwise-Hebrew PDF
    decode to garbage Latin-lookalike characters, which dilutes the
    text-layer ratio without the page actually containing less Hebrew).
    Returns None if OSD isn't usable at all (Tesseract/osd.traineddata
    missing, or nothing to render).
    """
    if not _tesseract_available():
        return None
    suffix = Path(file_path).suffix.lower()
    try:
        if suffix == ".pdf":
            images = _render_pdf_pages(file_path, dpi=150)  # OSD doesn't need full OCR DPI
            sample_img = images[0] if images else None
        elif suffix in _OCR_INPUT_EXTS:
            sample_img = Image.open(file_path)
        else:
            return None
        if sample_img is None:
            return None
        osd = pytesseract.image_to_osd(sample_img, config="--psm 0", output_type=Output.DICT)
        return (osd.get("script") or "").lower() == "hebrew"
    except Exception:
        # OSD needs osd.traineddata, which isn't guaranteed present on every
        # install -- fail open (None) rather than raise, so a missing OSD
        # pack degrades to "trust the text-layer signal", not a broken request.
        return None


def _detect_script_is_hebrew(file_path: str) -> bool | None:
    """
    Best-effort automatic Hebrew detection, so a Hebrew document still gets
    routed through the Tesseract pipeline even if the caller forgot to pass
    hebrew=True. This is not a hypothetical: a native-text-layer Hebrew PDF
    run through the default RapidOCR pipeline (which has no Hebrew model
    at all -- see the module docstring) produces text that's part-correct,
    part nonsense Latin-lookalike guesses. Every UI tab except Translate
    requires the user to manually tick a "document is in Hebrew" checkbox,
    so this is the actual failure mode to defend against, not an edge case.

    Combines two independent signals with OR, not "prefer one, fall back to
    the other": a confident text-layer ratio (_sample_text_hebrew_ratio) is
    trusted immediately since it's cheap, but a low/inconclusive ratio does
    NOT rule Hebrew out on its own -- it can just mean the PDF's embedded
    text layer is partially broken (see _detect_script_is_hebrew_from_image's
    docstring). So the image-based OSD check still runs and can independently
    confirm Hebrew even when the text-layer signal alone would have said no.
    Given the asymmetric cost here (an unnecessary Hebrew-OCR pass on a
    non-Hebrew doc is cheap; missing a genuinely Hebrew doc produces
    unreadable output), this deliberately errs toward "yes".
    """
    suffix = Path(file_path).suffix.lower()
    if suffix not in _OCR_INPUT_EXTS:
        return None

    text_ratio = _sample_text_hebrew_ratio(file_path) if suffix == ".pdf" else None
    print(f"[document] {Path(file_path).name}: text-layer Hebrew ratio = {text_ratio}")
    if text_ratio is not None and text_ratio >= _HEBREW_SCRIPT_RATIO_THRESHOLD:
        return True

    image_signal = _detect_script_is_hebrew_from_image(file_path)
    print(f"[document] {Path(file_path).name}: image OSD Hebrew signal = {image_signal}")
    if image_signal is not None:
        return image_signal

    # No usable image signal (Tesseract/OSD unavailable) -- fall back to
    # whatever the text layer suggested, even if below threshold, rather
    # than returning an unhelpful None past this point.
    if text_ratio is not None:
        return text_ratio >= _HEBREW_SCRIPT_RATIO_THRESHOLD
    return None


def resolve_hebrew_flag(file_path: str, hebrew: bool) -> bool:
    """
    Returns the effective Hebrew flag for this file: the caller's explicit
    request, OR'd with best-effort auto-detection -- so a Hebrew document
    is routed correctly even if the UI checkbox for it was left unchecked.
    An explicit hebrew=True is always trusted outright and skips detection
    entirely. Exposed separately (not just inlined into convert_to_markdown)
    so callers like main.py can report back which OCR engine actually ran.
    """
    if hebrew:
        print(f"[document] {Path(file_path).name}: hebrew=True passed explicitly")
        return True
    if Path(file_path).suffix.lower() not in _OCR_INPUT_EXTS:
        return False
    resolved = bool(_detect_script_is_hebrew(file_path))
    print(f"[document] {Path(file_path).name}: resolved hebrew={resolved} (auto-detected)")
    return resolved


def convert_to_markdown(file_path: str, hebrew: bool = False) -> str:
    """
    Convert any supported document (PDF/DOCX/PPTX/HTML/image) to Markdown.
    Set hebrew=True to OCR through Tesseract directly instead of Docling's
    default RapidOCR pipeline -- only do this for documents actually in
    Hebrew. Only affects PDFs and raw images; native formats (DOCX/PPTX)
    never go through OCR at all, so the flag has no effect on them.

    hebrew=False is no longer taken purely at face value: resolve_hebrew_flag
    auto-detects Hebrew content and overrides it when needed (see that
    function's docstring for why this matters in practice).
    """
    hebrew = resolve_hebrew_flag(file_path, hebrew)
    if hebrew and Path(file_path).suffix.lower() in _OCR_INPUT_EXTS:
        return _ocr_hebrew(file_path)

    result = _default_converter.convert(Path(file_path))
    markdown_text = html.unescape(result.document.export_to_markdown())
    # Defensive: same bidi-control-character cleanup as the Hebrew/Tesseract
    # path, applied here too in case RapidOCR ever emits stray marks on
    # Arabic (or any other RTL-script) content.
    return strip_bidi_controls(markdown_text)


def _paragraph_is_rtl(text: str, threshold: float = 0.3) -> bool:
    """
    A paragraph is treated as RTL if a meaningful share of its *letters*
    are Hebrew/Arabic -- not "contains any RTL character at all". That
    keeps a mostly-English paragraph that happens to mention one Hebrew
    name from getting force-flipped to right-aligned.
    """
    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return False
    rtl_letters = sum(1 for ch in letters if 0x0590 <= ord(ch) <= 0x06FF)
    return (rtl_letters / len(letters)) >= threshold


def _set_paragraph_rtl(paragraph) -> None:
    """
    Marks a python-docx paragraph as right-to-left and right-aligned.

    pandoc's markdown->docx conversion has no concept of paragraph
    direction -- every paragraph comes out as an ordinary left-to-right
    Word paragraph regardless of script. Word's renderer still reorders
    individual Hebrew/Arabic glyphs correctly (that's the Unicode bidi
    algorithm, and it happens automatically), but without the paragraph's
    w:bidi flag set, Word (a) left-aligns the whole paragraph, which reads
    backwards to an RTL reader, and (b) can misplace weakly-directional
    runs -- numbers, dates, embedded Latin words -- relative to the
    surrounding Hebrew/Arabic text. That second failure mode is the same
    root problem ui.py's _anchor_rtl_lines works around for the Gradio
    textbox; here we set the real OOXML property instead, since a Word
    document actually has one.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        rPr.append(rtl)


def _apply_rtl_formatting(docx_path: str) -> None:
    """
    Walks every paragraph in the generated docx -- including inside tables,
    since invoices/receipts routinely have RTL table content -- and flips
    direction on any paragraph that's predominantly Hebrew/Arabic.
    """
    doc = DocxDocument(docx_path)
    changed = False

    def _process(paragraphs):
        nonlocal changed
        for p in paragraphs:
            if _paragraph_is_rtl(p.text):
                _set_paragraph_rtl(p)
                changed = True

    _process(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _process(cell.paragraphs)

    if changed:
        doc.save(docx_path)


def export_docx(markdown_text: str, output_path: str) -> str:
    """
    Convert Markdown text to a .docx file via pandoc.
    Used for: PDF -> Word, and OCR'd image -> Word.
    Returns the output_path for convenience.
    """
    pypandoc.convert_text(
        markdown_text,
        to="docx",
        format="md",
        outputfile=output_path,
    )
    # pandoc's output is direction-agnostic (see _apply_rtl_formatting's
    # docstring) -- fix up any Hebrew/Arabic paragraphs after the fact.
    _apply_rtl_formatting(output_path)
    return output_path


def convert_file_to_docx(input_path: str, output_path: str, hebrew: bool = False) -> str:
    """
    Full pipeline: any supported input (PDF, image, PPTX, HTML...) -> Word.
    This is what backs the PDF->Word and Image(OCR)->Word conversion endpoints.
    """
    markdown_text = convert_to_markdown(input_path, hebrew=hebrew)
    return export_docx(markdown_text, output_path)


# Matches a sentence boundary: one of .!? followed by whitespace, followed
# by a character that plausibly starts a new sentence.
#
# IMPORTANT: the lookahead character class used to be Latin-only
# (A-Z, 0-9, and accented Latin). That meant it NEVER matched inside
# Hebrew, Arabic, Cyrillic, or CJK text -- there's no such thing as an
# "uppercase" Hebrew letter, so a period-plus-Hebrew-letter never looked
# like a sentence boundary. The practical effect: every long Hebrew
# paragraph was chunked as a single oversized, unsplit block (observed:
# 100-180 word chunks against an intended ~400-char/60-90-word target),
# and translate.py's own sentence-retry fallback degraded to a no-op for
# the same reason (a "sentence list" of length 1 can't be retried
# piecewise). MADLAD-400-3B is documented (see translate.py) to be much
# less reliable on long multi-sentence blocks than short ones, so this
# was the primary cause of Hebrew-source translations failing on every
# single chunk. Fixed by recognizing common non-Latin scripts as valid
# sentence-starting characters too.
SENTENCE_SPLIT_RE = re.compile(
    r"(?<=[.!?])\s+(?=[A-Z0-9\u00c0-\u024f"  # Latin (original)
    r"\u0400-\u04ff"  # Cyrillic
    r"\u0590-\u05ff"  # Hebrew
    r"\u0600-\u06ff"  # Arabic
    r"\u4e00-\u9fff"  # CJK unified ideographs
    r"])"
)


def _split_into_sentences(paragraph: str) -> list[str]:
    """Best-effort sentence splitter. Not perfect (abbreviations, decimals),
    but good enough to keep chunks short -- which matters far more than
    perfect boundaries for translation reliability."""
    sentences = SENTENCE_SPLIT_RE.split(paragraph.strip())
    return [s.strip() for s in sentences if s.strip()]


def _hard_split(unit: str, max_chars: int) -> list[str]:
    """
    Force-splits an oversized unit on whitespace/word boundaries, greedily
    packing words up to max_chars.

    This is a last-resort safety net for chunk_text below: sentence
    splitting is now script-aware (see SENTENCE_SPLIT_RE) but can still
    hand back a single "sentence" that's still too long -- a genuine
    run-on sentence, text with no recognized sentence punctuation at all,
    or some future script/punctuation style this regex doesn't cover.
    Rather than silently letting an oversized chunk through again (which is
    exactly how the Hebrew bug above went unnoticed), every unit is now
    guaranteed to fit within max_chars by the time it reaches the model.
    """
    if len(unit) <= max_chars:
        return [unit]
    words = unit.split()
    if not words:
        return [unit]
    pieces, current = [], ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) > max_chars and current:
            pieces.append(current)
            current = word
        else:
            current = candidate
    if current:
        pieces.append(current)
    return pieces


def chunk_text(markdown_text: str, max_chars: int = 400) -> list[str]:
    """
    Paragraph- and sentence-aware chunking so translation chunks stay small.

    max_chars was lowered from 800 to 400: MADLAD-400-3B (int8) is noticeably
    less reliable translating long, multi-sentence blocks in one shot than
    short ones -- observed failure mode is the model translating correctly
    for a while, then degenerating into a repetition loop and drifting back
    into the source language partway through. Paragraphs longer than
    max_chars are now split into individual sentences (not just left as one
    oversized chunk), and sentences are packed back together up to the limit
    so short sentences still get batched efficiently.

    Sentence splitting is script-aware (see SENTENCE_SPLIT_RE) and, as a
    final safety net, _hard_split guarantees no single piece handed to the
    packer ever exceeds max_chars, regardless of script or punctuation.
    """
    paragraphs = [p for p in markdown_text.split("\n\n") if p.strip()]
    chunks, current = [], ""

    def flush():
        nonlocal current
        if current.strip():
            chunks.append(current.strip())
        current = ""

    for p in paragraphs:
        units = [p] if len(p) <= max_chars else _split_into_sentences(p)
        for unit in units:
            for piece in _hard_split(unit, max_chars):
                if len(current) + len(piece) > max_chars and current:
                    flush()
                current += piece + " "

    flush()
    return chunks